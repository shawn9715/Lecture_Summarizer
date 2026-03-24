에러가 무사히 해결되었다니 정말 다행입니다! 드디어 지긋지긋한 환경 설정의 늪에서 완전히 빠져나오셨네요. 🎉

올려주신 통코드를 꼼꼼히 리뷰해 보았습니다. 코딩 파트너로서 아주 기쁜 소식을 전해드리자면, 회원님이 올려주신 코드 안에 이미 yt-dlp를 활용한 투트랙(웹 자동 추출 vs 파일 업로드) 로직이 완벽하게 들어가 있었습니다! (제가 어제 짜드린 코드를 아주 정확하게 반영해 주셨어요. 👍)

다만, 회원님의 예리한 지적대로 **"사용자가 화면(UI)만 봤을 때는 어떤 링크를 넣어야 하는지 헷갈릴 수 있다"**는 점이 핵심입니다.

따라서, 사용자가 직관적으로 이해할 수 있도록 UI를 **'쓰리트랙(3-Track)'**으로 확실하게 쪼개고, 그에 맞춰 다운로드 방식도 최적화하여 안정성을 극한으로 끌어올리겠습니다.

💡 파트너의 보강 전략 (UI 및 안정성 최적화)
🌐 1. 일반 웹페이지 주소: (yt-dlp 작동) 페이지 안의 플레이어를 AI가 샅샅이 뒤져서 영상을 뽑아옵니다.

🔗 2. 다이렉트 MP4 링크: (requests 작동) 군더더기 없이 MP4 원본 링크만 빠르게 쏙 다운로드합니다.

📁 3. 직접 파일 업로드: 보안이 너무 철저한 인강 사이트용 최후의 보루입니다.

자, 기존 aid.py 코드에서 아래의 새로운 전체 코드로 덮어씌워 주세요! (UI가 아주 예쁘고 명확하게 바뀌었을 겁니다.)

Python
import streamlit as st
import fitz
import google.generativeai as genai
import os
import time
import requests
import tempfile
import io
import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import pymongo
import yt_dlp

st.set_page_config(page_title="사회복지사 강의 요약기 PRO", page_icon="📚", layout="wide")
st.title("📚 사회복지사 강의 요약 웹서비스 (하이브리드 모드)")

# ---------------------------------------------------------
# 1. 라이브러리 폴더 준비 (플랜 B용)
# ---------------------------------------------------------
os.makedirs("library/pdfs", exist_ok=True)
os.makedirs("library/videos", exist_ok=True)

# ---------------------------------------------------------
# 2. 하이브리드 데이터베이스 연결 (플랜 A -> 플랜 B)
# ---------------------------------------------------------
use_db = False
pdf_collection = None
video_collection = None

@st.cache_resource
def init_db_connection():
    client = pymongo.MongoClient(st.secrets["MONGO_URI"], serverSelectionTimeoutMS=3000)
    client.server_info() 
    return client

try:
    if "MONGO_URI" in st.secrets and st.secrets["MONGO_URI"]:
        db_client = init_db_connection()
        db = db_client["LectureDB"]        
        pdf_collection = db["pdfs"]        
        video_collection = db["videos"]
        use_db = True
except Exception as e:
    use_db = False
    db_error_msg = str(e)

# ---------------------------------------------------------
# 3. 모델 및 워드 변환 함수
# ---------------------------------------------------------
def get_optimal_model():
    available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
    for name in available_models:
        if 'gemini-1.5-pro' in name: return name
    for name in available_models:
        if 'gemini-1.5-flash' in name: return name
    for name in available_models:
        if 'gemini' in name: return name
    raise ValueError("사용 가능한 모델을 찾을 수 없습니다.")

def generate_word_file(title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    lines = content.split('\n')
    first_subheading_done = False
    for line in lines:
        line = line.strip()
        if not line: continue
        is_subheading = bool(re.match(r'^\d+\.', line))
        if is_subheading:
            if first_subheading_done: doc.add_paragraph()
            first_subheading_done = True
            p = doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            
        parts = re.split(r'(\[강조\].*?\[/강조\])', line)
        for part in parts:
            if part.startswith('[강조]') and part.endswith('[/강조]'):
                run = p.add_run(part[4:-5])
                if is_subheading: run.bold = True
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                if part:
                    run = p.add_run(part)
                    if is_subheading: run.bold = True
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# 4. 사이드바 UI 구성
# ---------------------------------------------------------
saved_api_key = st.secrets["GEMINI_API_KEY"] if "GEMINI_API_KEY" in st.secrets else ""
with st.sidebar:
    st.header("⚙️ 기본 설정")
    api_key = st.text_input("Gemini API Key", value=saved_api_key, type="password")
    if api_key:
        genai.configure(api_key=api_key)
        
    st.divider()
    st.subheader("🗄️ 저장소 상태")
    if use_db:
        st.success("🟢 클라우드 DB 연결됨 (영구 저장 모드)")
    else:
        st.warning("🟡 로컬 임시 저장 모드 작동 중")
        st.caption("DB 연결에 실패하여 서버 내 임시 폴더를 사용합니다.")

tab1, tab2 = st.tabs(["📂 1단계: 교안 라이브러리 등록", "🎬 2단계: 동영상 분석 및 요약"])

# ---------------------------------------------------------
# [탭 1] 교안 PDF 등록
# ---------------------------------------------------------
with tab1:
    st.header("1. 교안 PDF 등록하기")
    course_name_input = st.text_input("🏷️ 강의 이름 입력 (예: 사회복지학개론)")
    pdf_file = st.file_uploader("📄 교안 PDF 파일 업로드", type=['pdf'])
    
    if st.button("💾 교안 저장하기"):
        if not course_name_input or not pdf_file:
            st.warning("강의 이름과 PDF 파일을 모두 입력해 주세요.")
        else:
            with st.spinner("PDF 문서 텍스트를 추출하여 저장 중입니다..."):
                doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
                full_text = "".join([page.get_text() for page in doc])
                
                if use_db:
                    pdf_collection.update_one(
                        {"course_name": course_name_input},
                        {"$set": {"content": full_text}},
                        upsert=True
                    )
                    st.success(f"🎉 '{course_name_input}' 교안이 클라우드 DB에 영구 저장되었습니다!")
                else:
                    with open(f"library/pdfs/{course_name_input}.txt", "w", encoding="utf-8") as f:
                        f.write(full_text)
                    st.success(f"⚠️ '{course_name_input}' 교안이 로컬 폴더에 임시 저장되었습니다!")

# ---------------------------------------------------------
# [탭 2] 동영상 분석 및 최종 요약 (🌟 3-Track UI 보강)
# ---------------------------------------------------------
with tab2:
    st.header("2. 동영상 분석 및 최종 요약본 생성")
    
    saved_pdfs = []
    if use_db:
        saved_pdfs_cursor = pdf_collection.find({}, {"course_name": 1, "_id": 0})
        saved_pdfs = [doc["course_name"] for doc in saved_pdfs_cursor]
    else:
        saved_pdfs = [f.replace(".txt", "") for f in os.listdir("library/pdfs") if f.endswith(".txt")]
    
    if not saved_pdfs:
        st.info("👈 먼저 '1단계' 탭에서 교안을 등록해 주세요.")
    else:
        selected_course = st.selectbox("📚 분석할 강의 선택", saved_pdfs)
        col1, col2 = st.columns(2)
        with col1: week = st.number_input("주차 (예: 1)", min_value=1, step=1)
        with col2: session = st.number_input("강 (예: 1)", min_value=1, step=1)
            
        st.divider()
        st.subheader("🎬 동영상 입력 방식 선택")
        
        # 🌟 3가지 방식으로 명확하게 나눈 UI
        input_method = st.radio(
            "어떤 방식으로 동영상을 분석할까요?", 
            [
                "🌐 1. 일반 웹페이지 주소 (동영상 자동 추출)", 
                "🔗 2. 다이렉트 MP4 링크 (https://...mp4)", 
                "📁 3. 동영상 파일 직접 업로드 (가장 확실한 방법)"
            ]
        )
        
        video_url = ""
        video_file = None
        
        if "일반 웹페이지" in input_method:
            st.info("💡 사이트 내에 숨겨진 동영상 플레이어를 자동으로 찾아냅니다.")
            video_url = st.text_input("🌐 웹페이지 주소 입력 (예: https://...)")
        elif "다이렉트 MP4" in input_method:
            st.info("💡 동영상 원본 링크(.mp4)를 직접 입력하여 빠르게 다운로드합니다.")
            video_url = st.text_input("🔗 MP4 링크 입력 (예: https://...mp4)")
        else:
            st.info("💡 유료 인강처럼 로그인이 필요한 사이트는 보안상 직접 업로드해야 합니다.")
            video_file = st.file_uploader("📁 동영상 파일 업로드", type=['mp4', 'avi', 'mov'])
        
        if st.button("🚀 분석 및 워드(Word) 요약본 생성 시작"):
            if not api_key:
                st.error("사이드바에 API 키를 입력해 주세요.")
            elif ("웹페이지" in input_method or "MP4" in input_method) and (not video_url):
                st.warning("링크 주소를 입력해 주세요.")
            elif ("업로드" in input_method) and (not video_file):
                st.warning("동영상 파일을 업로드해 주세요.")
            else:
                try:
                    tmp_video = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
                    temp_video_path = tmp_video.name
                    tmp_video.close()

                    # 🌟 선택한 방식에 따라 다른 다운로드 엔진 사용
                    with st.spinner("1/4. 영상을 준비하고 있습니다... ⏳"):
                        if "일반 웹페이지" in input_method:
                            # 플랜 A: yt-dlp로 플레이어 추출
                            ydl_opts = {'format': 'best', 'outtmpl': temp_video_path, 'quiet': True, 'no_warnings': True}
                            try:
                                with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                                    ydl.download([video_url])
                            except Exception as e:
                                raise Exception(f"보안 설정으로 인해 영상 추출에 실패했습니다. 파일 직접 업로드 방식을 사용해 주세요. (에러: {e})")
                        elif "다이렉트 MP4" in input_method:
                            # 플랜 B: requests로 다이렉트 다운로드 (가볍고 빠름)
                            try:
                                response = requests.get(video_url, stream=True)
                                response.raise_for_status()
                                with open(temp_video_path, 'wb') as f:
                                    for chunk in response.iter_content(chunk_size=8192):
                                        f.write(chunk)
                            except Exception as e:
                                raise Exception(f"MP4 링크 다운로드에 실패했습니다. 주소가 정확한지 확인해 주세요. (에러: {e})")
                        else:
                            # 플랜 C: 로컬 파일 저장
                            with open(temp_video_path, 'wb') as f:
                                f.write(video_file.read())

                    st.write("2/4. AI 서버에 영상을 올리고 처리하는 중입니다. (수 분 소요) 🧠")
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    optimal_model_name = get_optimal_model()
                    model = genai.GenerativeModel(optimal_model_name)
                    video_upload = genai.upload_file(path=temp_video_path)
                    
                    start_time = time.time()
                    while True:
                        file_info = genai.get_file(video_upload.name)
                        state = file_info.state.name
                        elapsed = int(time.time() - start_time)
                        if state == "PROCESSING":
                            status_text.text(f"⏳ 서버에서 영상 분석 준비 중... (현재 {elapsed}초 경과)")
                            progress_bar.progress(min(elapsed, 95))
                            time.sleep(5)
                        elif state == "ACTIVE":
                            progress_bar.progress(100)
                            status_text.success(f"✅ AI 영상 분석 준비 완료! (총 {elapsed}초 소요)")
                            break
                        elif state == "FAILED":
                            raise Exception("AI 서버에서 동영상 처리에 실패했습니다.")
                        else:
                            time.sleep(5)

                    with st.spinner("3/4. 영상의 핵심 내용을 추출하고 있습니다... ✍️"):
                        video_prompt = """
                        이 강의 영상에서 다음 사항을 추출해 주세요:
                        1. 강사가 음성으로 "시험에 나온다", "중요하다" 등으로 특별히 강조한 내용
                        2. 화면 상에서 하이라이트(밑줄, 별표, 빨간 글씨 등) 처리된 핵심 내용
                        """
                        video_response = model.generate_content([video_upload, video_prompt])
                        video_key_points = video_response.text
                        
                        genai.delete_file(video_upload.name)
                        os.unlink(temp_video_path)
                        
                        if use_db:
                            video_collection.update_one(
                                {"course_name": selected_course, "week": week, "session": session},
                                {"$set": {"content": video_key_points}},
                                upsert=True
                            )
                        else:
                            video_lib_path = f"library/videos/{selected_course}_{week}주차_{session}강_핵심내용.txt"
                            with open(video_lib_path, "w", encoding="utf-8") as f:
                                f.write(video_key_points)

                    with st.spinner("4/4. 최종 노트 필기본을 작성 중입니다... 📝"):
                        if use_db:
                            course_data = pdf_collection.find_one({"course_name": selected_course})
                            course_full_text = course_data["content"]
                        else:
                            with open(f"library/pdfs/{selected_course}.txt", "r", encoding="utf-8") as f:
                                course_full_text = f.read()
                            
                        final_prompt = f"""
                        당신은 훌륭한 사회복지사 학습 조교입니다.
                        아래 두 자료를 대조하여 '{selected_course}'의 {week}주차 {session}강 노트 필기본을 작성해 주세요.

                        [자료 1: 영상 핵심 내용]
                        {video_key_points}

                        [자료 2: 전체 교안 텍스트]
                        {course_full_text}

                        [작성 지시 사항 - 매우 중요]
                        1. 내용 대조: [자료 1]의 핵심 내용을 바탕으로 [자료 2]에서 상세한 설명을 찾아 요약하세요.
                        2. 공통 강조점 하이라이트: 양쪽 모두에서 중요하게 다뤄진 단어나 문장은 앞뒤에 [강조] 와 [/강조] 태그를 붙이세요. 
                        3. 소제목 규칙: 내용의 주제가 바뀔 때는 반드시 "1. 주제명" 처럼 숫자와 마침표로 시작하는 소제목을 적어주세요.
                        4. 포맷: 마크다운 기호(**, # 등)는 쓰지 말고, 번호와 기호(-, ※)만 사용하여 어르신들이 읽기 편한 일반 텍스트로 정리해 주세요.
                        """
                        final_response = model.generate_content(final_prompt)
                        final_summary = final_response.text

                    st.success("🎉 모든 분석 및 정리가 완료되었습니다!")
                    st.divider()
                    
                    final_title = f"[{selected_course}] - {week}주차 {session}강 요약 노트"
                    preview_text = final_summary.replace("[강조]", "").replace("[/강조]", "")
                    st.subheader("👀 요약본 미리보기")
                    st.text(preview_text)
                    st.divider()
                    
                    word_buffer = generate_word_file(final_title, final_summary)
                    st.download_button(
                        label="📄 굵은 글씨 및 형광펜이 적용된 워드 파일 다운로드",
                        data=word_buffer,
                        file_name=f"{final_title}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except requests.exceptions.RequestException as e:
                    st.error(f"인터넷 연결 에러: {e}")
                except Exception as e:
                    st.error(f"오류가 발생했습니다: {e}")
